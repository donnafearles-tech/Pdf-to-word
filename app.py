import os
import re
import time
import docx
import streamlit as st
from groq import Groq

# =====================================================================
# IMPORTACIONES OFICIALES DEL SDK DE ADOBE (V4)
# =====================================================================
from adobe.pdfservices.operation.auth.service_principal_credentials import ServicePrincipalCredentials
from adobe.pdfservices.operation.pdf_services import PDFServices
from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
from adobe.pdfservices.operation.pdfjobs.jobs.export_pdf_job import ExportPDFJob
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_params import ExportPDFParams
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_target_format import ExportPDFTargetFormat
from adobe.pdfservices.operation.pdfjobs.result.export_pdf_result import ExportPDFResult

# =====================================================================
# CONFIGURACIÓN DE LA PÁGINA DE STREAMLIT
# =====================================================================
st.set_page_config(
    page_title="Conversor Editorial PDF", 
    page_icon="📚", 
    layout="centered"
)

# =====================================================================
# 1. MOTOR DE CONVERSIÓN (ADOBE SDK V4) - BLINDADO
# =====================================================================
def convertir_pdf_a_word_adobe(input_pdf_path, output_docx_path, client_id, client_secret):
    """
    Convierte un PDF a DOCX usando la API oficial de Adobe (SDK v4).
    Maneja correctamente la lectura y escritura de bytes puros.
    """
    try:
        credentials = ServicePrincipalCredentials(
            client_id=client_id, 
            client_secret=client_secret
        )
        pdf_services = PDFServices(credentials=credentials)

        with open(input_pdf_path, 'rb') as f:
            pdf_bytes = f.read()
            
        asset = pdf_services.upload(input_stream=pdf_bytes, mime_type=PDFServicesMediaType.PDF)

        params = ExportPDFParams(target_format=ExportPDFTargetFormat.DOCX)
        job = ExportPDFJob(input_asset=asset, export_pdf_params=params)
        
        location = pdf_services.submit(job)
        pdf_services_response = pdf_services.get_job_result(location, ExportPDFResult)
        
        result_asset = pdf_services_response.get_result().get_asset()
        stream_asset = pdf_services.get_content(result_asset)

        with open(output_docx_path, "wb") as f:
            f.write(stream_asset.get_input_stream())
            
        return True

    except Exception as e:
        st.error(f"Error fatal en Adobe PDF Services al convertir: {str(e)}")
        return False

# =====================================================================
# 2. HEURÍSTICAS DE FILTRADO Y LIMPIEZA (TEXTO Y GRÁFICOS)
# =====================================================================
def limpiar_imagenes_pequenas(doc, min_width_cm=1.5, min_height_cm=1.5):
    """
    Itera sobre las imágenes incrustadas y elimina las que sean más pequeñas 
    que el umbral especificado para purgar manchas, logos o ruido de escaneo.
    """
    imagenes_eliminadas = 0
    for shape in doc.inline_shapes:
        try:
            ancho = shape.width.cm
            alto = shape.height.cm
            
            if ancho < min_width_cm or alto < min_height_cm:
                nodo_imagen = shape._inline
                nodo_imagen.getparent().remove(nodo_imagen)
                imagenes_eliminadas += 1
        except Exception:
            continue
            
    return imagenes_eliminadas

def pre_limpiar_ocr(texto):
    """
    Conserva únicamente el alfabeto inglés/español, números y puntuación estándar.
    Elimina ráfagas de símbolos basura del OCR antes de procesar con la IA.
    """
    # Expresión regular inclusiva (filtra todo lo que NO sea letra es/en, número o puntuación básica)
    patron_permitido = r'[^a-zA-ZáéíóúÁÉÍÓÚñÑüÜ0-9\s.,;:\-!?¿¡"\'\(\)\[\]/]'
    texto_limpio = re.sub(patron_permitido, '', texto)
    
    # Colapsar espacios múltiples y saltos de línea huérfanos
    return re.sub(r'\s+', ' ', texto_limpio).strip()

# =====================================================================
# 3. MOTOR DE LIMPIEZA Y TRADUCCIÓN (GROQ) - CON RESILIENCIA DE CUOTA
# =====================================================================
def limpiar_y_traducir_con_groq(texto, groq_api_key, max_reintentos=3):
    """Llama a Groq para limpiar y traducir, blindado contra Rate Limits (429)."""
    cliente = Groq(api_key=groq_api_key)
    prompt_sistema = (
        "Eres un editor editorial experto en restauración de textos escaneados.\n"
        "1. Traduce el texto al ESPAÑOL de forma natural.\n"
        "2. Elimina basura de escaneo: símbolos sin sentido o sílabas rotas.\n"
        "3. Corrige la ortografía y puntuación.\n"
        "4. Devuelve ÚNICAMENTE el texto traducido. Sin introducciones."
    )
    
    for intento in range(max_reintentos):
        try:
            respuesta = cliente.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {"role": "system", "content": prompt_sistema},
                    {"role": "user", "content": texto}
                ],
                temperature=0.1,
                max_tokens=1500
            )
            return respuesta.choices[0].message.content.strip()
            
        except Exception as e:
            error_msg = str(e).lower()
            # Si el error es por saturación de tokens por minuto (Rate limit)
            if "rate limit" in error_msg or "429" in error_msg:
                st.warning(f"⏳ Cuota de Groq saturada. Esperando 60 segundos... (Intento {intento + 1}/{max_reintentos})")
                time.sleep(60)
            else:
                st.warning(f"Aviso: Un párrafo falló por red. Se mantendrá el original. Error: {str(e)}")
                return texto
                
    return texto

def procesar_docx_con_groq(docx_path, groq_api_key):
    """Itera sobre el Word aplicando pausas anti-baneo y autoguardado de seguridad."""
    doc = docx.Document(docx_path)
    
    texto_estado = st.empty()
    texto_estado.text("Fase 2a: Purgando imágenes minúsculas y ruido visual de escaneo...")
    
    # 1. Purgar imágenes inútiles primero
    img_eliminadas = limpiar_imagenes_pequenas(doc, min_width_cm=1.5, min_height_cm=1.5)
    st.info(f"🧹 Se eliminaron {img_eliminadas} artefactos visuales/imágenes pequeñas.")
    
    # 2. Inicializar bucle de traducción y corrección
    barra_progreso = st.progress(0)
    total_parrafos = len(doc.paragraphs)
    parrafos_procesados = 0
    
    for i, parrafo in enumerate(doc.paragraphs):
        texto_original = parrafo.text.strip()
        
        progreso = int(((i + 1) / total_parrafos) * 100)
        barra_progreso.progress(progreso)
        texto_estado.text(f"Limpiando y traduciendo párrafo {i + 1} de {total_parrafos}...")
        
        if not texto_original or texto_original.isdigit():
            continue
            
        texto_pre_limpio = pre_limpiar_ocr(texto_original)
        
        if len(texto_pre_limpio) > 3:
            texto_final = limpiar_y_traducir_con_groq(texto_pre_limpio, groq_api_key)
            
            estilo_previo = None
            if parrafo.runs and parrafo.runs[0].style:
                estilo_previo = parrafo.runs[0].style

            for run in parrafo.runs:
                run.text = ""
                
            nuevo_run = parrafo.add_run(texto_final)
            if estilo_previo:
                nuevo_run.style = estilo_previo
            
            parrafos_procesados += 1
            time.sleep(2.5)  # Pausa obligatoria para mitigar el límite TPM
            
        # 3. GUARDADO INCREMENTAL: Evita perder el progreso en ejecuciones largas
        if i > 0 and i % 50 == 0:
            doc.save(docx_path)
            texto_estado.text(f"💾 Guardado de seguridad incremental realizado en el párrafo {i}...")

    doc.save(docx_path)
    texto_estado.text(f"✅ Completado. {parrafos_procesados} párrafos mejorados.")
    barra_progreso.empty()

# =====================================================================
# 4. INTERFAZ DE USUARIO Y CONTROL DE FLUJO PRINCIPAL
# =====================================================================
st.title("Conversor Editorial: PDF a Word Limpio")
st.markdown("Sube tus archivos **PDF escaneados** para convertirlos a **Word**, traducirlos al español y remover ruido de OCR.")

try:
    ADOBE_CLIENT_ID = st.secrets["PDF_SERVICES_CLIENT_ID"]
    ADOBE_CLIENT_SECRET = st.secrets["PDF_SERVICES_CLIENT_SECRET"]
    GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
except KeyError as e:
    st.error(f"❌ Error crítico: Falta la credencial {e} en los Secrets de Streamlit.")
    st.stop()

archivo_subido = st.file_uploader("Selecciona el libro o documento en formato PDF", type=["pdf"])

if archivo_subido:
    if st.button("Comenzar Procesamiento Editorial", type="primary"):
        
        id_unico = str(int(time.time()))
        temp_pdf = f"temp_input_{id_unico}.pdf"
        temp_docx = f"temp_output_{id_unico}.docx"
        
        try:
            with open(temp_pdf, "wb") as f:
                f.write(archivo_subido.getbuffer())
                
            with st.spinner("Fase 1/2: Convirtiendo estructura del PDF a Word en servidores de Adobe..."):
                exito_adobe = convertir_pdf_a_word_adobe(
                    temp_pdf, temp_docx, ADOBE_CLIENT_ID, ADOBE_CLIENT_SECRET
                )
                
            if exito_adobe:
                with st.spinner("Fase 2/2: Inicializando Inteligencia Artificial para limpieza..."):
                    procesar_docx_con_groq(temp_docx, GROQ_API_KEY)
                    
                st.success("🎉 ¡El documento ha sido procesado y restaurado con éxito!")
                st.balloons()
                
                with open(temp_docx, "rb") as f:
                    st.download_button(
                        label="📥 Descargar Documento Word Limpio",
                        data=f,
                        file_name="Libro_Procesado_Limpio.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.error("❌ El proceso se detuvo porque la conversión de Adobe falló.")
                
        except Exception as e:
            st.error(f"Ha ocurrido un error inesperado en la aplicación: {str(e)}")
            
        finally:
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
